"""
Render the integrity test markdown report to PDF via python-docx → soffice.

Approach (per the markdown-to-pdf skill, "academic-paper-versioning" pitfall):
- Parse the markdown manually into a python-docx Document
- Use Times New Roman 10.5pt body, single column, Letter size, 0.85in margins
- Save .docx → convert to .pdf via soffice headless
- No headers/footers, no page numbers, no emoji
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path("/Users/jkm/Projects/dtu_full_code 2/reports/MODEL_INTEGRITY_TEST_REPORT.md")
DOCX_OUT = Path("/Users/jkm/Projects/dtu_full_code 2/reports/MODEL_INTEGRITY_TEST_REPORT.docx")
PDF_OUT = Path("/Users/jkm/Projects/dtu_full_code 2/reports/MODEL_INTEGRITY_TEST_REPORT.pdf")

text = SRC.read_text(encoding="utf-8")
lines = text.splitlines()

# Title block extraction
title = None
subtitle = None
i = 0
while i < len(lines):
    s = lines[i].strip()
    if title is None and s.startswith("# "):
        title = s[2:].strip()
        i += 1
        continue
    if title is not None and subtitle is None and s.startswith("## "):
        subtitle = s[3:].strip()
        i += 1
        continue
    if title is not None and s == "":
        i += 1
        continue
    if title is not None:
        break
    i += 1

body_lines = lines[i:]

# ---- Document setup ----
doc = Document()

# Letter size, 0.85" margins
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    # No headers/footers — leave default which is empty

# Default body style: Times New Roman 10.5pt, justify
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10.5)
# Set East Asian / complex script fonts too so it's consistent everywhere
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:ascii"), "Times New Roman")
rFonts.set(qn("w:hAnsi"), "Times New Roman")
rFonts.set(qn("w:cs"), "Times New Roman")
rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = style.paragraph_format
pf.space_after = Pt(3)
pf.line_spacing = 1.2


def add_inline_runs(paragraph, text):
    """Parse inline **bold** / *italic* / `code` and append runs to paragraph."""
    # Tokenize: split on **...**, *...*, `...`
    pattern = re.compile(r"(\*\*[^\*]+?\*\*)|(\*[^\*]+?\*)|(`[^`]+?`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("*"):
            r = paragraph.add_run(tok[1:-1])
            r.italic = True
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_heading(text, level):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(15)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(12)
    elif level == 3:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
    elif level == 4:
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p


def add_paragraph_text(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_runs(p, text)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline_runs(p, text)
    return p


def add_numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline_runs(p, text)
    return p


def add_hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(rows):
    """Add a table. First row is header (bold + light grey fill)."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Apply borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "666666")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= n_cols:
                continue
            cell = table.rows[ri].cells[ci]
            # Clear default empty paragraph
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
            # Header row: light grey shading
            if ri == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "E8E8E8")
                tcPr.append(shd)
    # Spacer paragraph after table
    doc.add_paragraph()


# ---- Title block ----
if title:
    add_heading(title, 1)
if subtitle:
    add_heading(subtitle, 2)
add_hr()

# ---- Body ----
i = 0
while i < len(body_lines):
    ln = body_lines[i].rstrip()

    # Headings
    m = re.match(r"^(#{1,6})\s+(.+)$", ln)
    if m:
        level = len(m.group(1))
        text_h = m.group(2).strip()
        # Word uses levels 1-9; map ours
        if level == 1:
            add_heading(text_h, 3)
        elif level == 2:
            add_heading(text_h, 4)
        else:
            add_heading(text_h, 4)
        i += 1
        continue

    # Horizontal rule
    if ln.strip() in ("---", "***"):
        add_hr()
        i += 1
        continue

    # Table: collect contiguous pipe rows
    if ln.lstrip().startswith("|"):
        tbl_rows = []
        while i < len(body_lines) and body_lines[i].lstrip().startswith("|"):
            row = body_lines[i].strip().strip("|")
            cells = [c.strip() for c in row.split("|")]
            # Skip separator row like |---|---|
            if all(re.match(r"^:?-+:?$", c) for c in cells):
                i += 1
                continue
            tbl_rows.append(cells)
            i += 1
        add_table(tbl_rows)
        continue

    # Bullet
    if re.match(r"^\s*-\s+", ln):
        text_b = re.sub(r"^\s*-\s+", "", ln)
        add_bullet(text_b)
        i += 1
        continue

    # Numbered
    if re.match(r"^\s*\d+\.\s+", ln):
        text_n = re.sub(r"^\s*\d+\.\s+", "", ln)
        add_numbered(text_n)
        i += 1
        continue

    # Blank line
    if not ln.strip():
        i += 1
        continue

    # Regular paragraph — collect consecutive non-blank lines
    para = ln
    i += 1
    while i < len(body_lines):
        nxt = body_lines[i].rstrip()
        if (not nxt.strip()
            or nxt.lstrip().startswith("|")
            or nxt.lstrip().startswith("#")
            or re.match(r"^\s*-\s+", nxt)
            or re.match(r"^\s*\d+\.\s+", nxt)
            or nxt.strip() in ("---", "***")):
            break
        para += " " + nxt.strip()
        i += 1
    add_paragraph_text(para)

# Save
doc.save(str(DOCX_OUT))
print(f"Wrote {DOCX_OUT} ({DOCX_OUT.stat().st_size:,} bytes)")

# Convert via soffice headless
import subprocess
r = subprocess.run(
    ["soffice", "--headless", "--convert-to", "pdf",
     "--outdir", str(PDF_OUT.parent),
     str(DOCX_OUT)],
    capture_output=True, text=True, timeout=180
)
if r.returncode != 0:
    print("soffice stderr:", r.stderr)
    raise SystemExit(1)
print(f"Wrote {PDF_OUT} ({PDF_OUT.stat().st_size:,} bytes)")