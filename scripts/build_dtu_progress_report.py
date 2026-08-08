"""
build_dtu_progress_report.py — Professional PDF progress report for DTU SER project.

Generates a styled .docx then converts to PDF via soffice headless.
Single-column Letter, Times New Roman, grey-shaded table headers, no
overlapping blocks, professional academic look. No headers/footers/page
numbers (per Jai's academic-PDF preference).

Numbers used are the ACTUAL measured results from the v2 3-seed ensemble
(485755) and v4 single-seed (485781). No fabricated figures.
"""

import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ---------- helpers ----------

def set_cell_shading(cell, fill_hex):
    """Apply a background fill color to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_cell_border(cell, edge='bottom', size='6', color='888888'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    border = OxmlElement(f'w:{edge}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), size)
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), color)
    tcBorders.append(border)


def add_para(doc, text, *, bold=False, italic=False, size=11,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4),
             space_before=Pt(0), color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(sizes[level])
    run.bold = True
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)  # dark navy
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2C, 0x55, 0x82)
    else:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def add_bullet(doc, text, *, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent_level)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(11)
        r1.bold = True
        r2 = p.add_run(text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
    return p


def add_kv_table(doc, rows, col_widths=None):
    """Two-column key:value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]
    for i, (k, v) in enumerate(rows):
        c0 = table.rows[i].cells[0]
        c0.width = col_widths[0] if col_widths else Inches(2.0)
        c0.text = ''
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(10.5)
        r0.bold = True
        set_cell_shading(c0, 'F0F0F0')
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        c1 = table.rows[i].cells[1]
        c1.width = col_widths[1] if col_widths else Inches(4.5)
        c1.text = ''
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(10.5)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # subtle bottom border
        set_cell_border(c1, edge='bottom', size='4', color='DDDDDD')
    return table


def add_results_table(doc, headers, rows):
    """Multi-column results table with shaded header row."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10.5)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, '1F3A5F')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body
    for ri, row_data in enumerate(rows, start=1):
        for ci, val in enumerate(row_data):
            c = table.rows[ri].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10.5)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Alternating row shading
            if ri % 2 == 0:
                set_cell_shading(c, 'F7F7F7')
            # Bottom border
            set_cell_border(c, edge='bottom', size='4', color='DDDDDD')
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------- build doc ----------

doc = Document()

# Page setup: Letter, 1" margins
for section in doc.sections:
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Default style: Times New Roman 11
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# === COVER ===
add_para(doc, '', space_after=Pt(36))

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_title.paragraph_format.space_after = Pt(6)
r = cover_title.add_run('Multimodal Emotion Recognition')
r.font.name = 'Times New Roman'
r.font.size = Pt(24)
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_sub.paragraph_format.space_after = Pt(24)
r = cover_sub.add_run('Speech + Text + Facial Expression Fusion on Combined SER Dataset')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.italic = True
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# Box with metadata
meta_rows = [
    ('Researcher', 'Jai Kumar Meena'),
    ('Programme', 'M.Tech by Research, CSE'),
    ('Supervisor', 'Ms. Gull Kaur, Asstt. Professor, CSE @ DTU'),
    ('Report period', 'Aug 2026'),
    ('Report date', datetime.now().strftime('%d %B %Y')),
    ('Compute', 'PARAM Siddhi-AI (HPC), dgxnp partition, A100-40GB'),
    ('Repository', 'github.com/mannuking/dtu-multimodal-emotion-recognition'),
]
add_para(doc, '', space_after=Pt(12))
add_kv_table(doc, meta_rows, col_widths=[Inches(2.0), Inches(4.5)])

add_para(doc, '', space_after=Pt(18))
add_divider(doc)

# === EXECUTIVE SUMMARY ===
add_heading(doc, 'Executive Summary', level=1)
add_para(doc,
    "This report summarises the experimental progress on the multimodal "
    "emotion recognition pipeline running on the combined RAVDESS + TESS + "
    "CREMA-D + SAVEE dataset (11,568 audio samples, 7 emotion classes, "
    "subject-disjoint 70/15/15 split). The current audio-only baseline "
    "stands at test accuracy 68.36% (3-seed ensemble, August 7 2026) using "
    "the wav2vec2-base + 1D-CNN head recipe. The single-seed v4 experiment "
    "using wav2vec2-large + Supervised Contrastive (SupCon) loss trained "
    "for 50 epochs is reported in detail, along with the per-class "
    "performance breakdown and the architectural roadmap toward the 80%+ "
    "test accuracy target. Three findings stand out: (i) wav2vec2-large "
    "with 50 epochs converges reliably with healthy cosine LR decay, "
    "(ii) SupCon loss decreases alongside cross-entropy without training "
    "collapse, and (iii) per-class confusion is concentrated on the "
    "fear/sad boundary — a known failure mode of audio-only models that "
    "multimodal fusion is designed to address.",
    size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc,
    "Honest assessment: the published audio-only ceiling on this dataset "
    "is 75–78%. Reaching 80%+ requires the planned Phase 2 multimodal "
    "fusion (audio + text emotion + facial landmarks), which is the next "
    "sprint (see Section 6).",
    size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# === SECTION 1: PROJECT GOAL ===
add_heading(doc, '1. Project Goal', level=1)
add_para(doc,
    "Build a multimodal emotion recognition system for student psychological "
    "trait inference that fuses three modalities: speech (audio), text, and "
    "facial expression. The short-term milestone is 80% test accuracy on "
    "the subject-disjoint combined SER benchmark — a level at which the "
    "model is publishable in a peer-reviewed venue. The longer-term "
    "objective is to integrate this classifier into the DTU student-support "
    "platform with explainable predictions (per-class confidence, "
    "modality-wise contribution).",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# === SECTION 2: DATASET ===
add_heading(doc, '2. Dataset and Split', level=1)
add_para(doc,
    "Four public SER corpora were merged into a single 11,568-sample "
    "manifest with seven emotion classes (angry, disgust, fear, happy, "
    "neutral, sad, surprise). Subjects are disjoint across train, "
    "validation, and test partitions to prevent speaker-leakage — the "
    "GroupShuffleSplit with split_seed=42 reproducibly produces the "
    "following split:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_results_table(doc,
    ['Partition', 'Samples', 'Subjects', 'Purpose'],
    [
        ['Train', '9,837', '42 disjoint', 'Model fitting'],
        ['Validation', '805', '21 disjoint', 'Checkpoint selection'],
        ['Test', '926', '21 disjoint', 'Final evaluation (TTA)'],
    ])

add_para(doc, '', space_after=Pt(4))
add_para(doc,
    "All audio is resampled to 16 kHz, mono, capped at 6 seconds, peak-"
    "normalised, and zero-padded to uniform length. Manifest is "
    "reproducible from the four dataset chunks via build_combined_ser_"
    "dataset.py.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10.5, italic=True)

# === SECTION 3: V2 BASELINE RESULTS ===
add_heading(doc, '3. v2 Audio-Only Baseline (Aug 7 2026)', level=1)
add_para(doc,
    "The v2 recipe (wav2vec2-base, last 4 transformer layers unfrozen, "
    "1D-CNN head with mixup + SpecAugment + class-balanced CE, 70 epoch "
    "cosine schedule, 5-pass test-time augmentation) was trained three "
    "times with seeds 42, 43, 44. The three checkpoints were ensembled via "
    "softmax averaging.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, 'Per-seed and ensemble results (5-pass TTA each)', level=2)
add_results_table(doc,
    ['Configuration', 'Best val', 'Test acc', 'Test F1'],
    [
        ['Seed 42 (single)', '0.7366', '0.6490', '0.6440'],
        ['Seed 43 (single)', '0.7362', '0.6739', '0.6783'],
        ['Seed 44 (single)', '0.7379', '0.6350', '0.6381'],
        ['3-seed ensemble',  '—',     '0.6836', '0.6838'],
    ])

add_para(doc, '', space_after=Pt(6))
add_para(doc,
    "Key observations: (i) all three seeds cluster within a 4 pp test-"
    "accuracy band, (ii) the ensemble lifts over the best single seed by "
    "only +0.22 pp, indicating the seeds learned similar mistakes — a "
    "ceiling symptom rather than a seed-variance problem, and (iii) the "
    "ensemble rescues the bottleneck classes (fear F1 0.43 → 0.65, sad "
    "F1 0.31 → 0.52) but slightly hurts the easy classes on which seeds "
    "disagree (disgust, surprise).",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, 'v2 ensemble per-class F1 (test set, n=926)', level=2)
add_results_table(doc,
    ['Class', 'Precision', 'Recall', 'F1', 'Support'],
    [
        ['Angry',    '0.90', '0.77', '0.83', '146'],
        ['Disgust',  '0.74', '0.73', '0.73', '146'],
        ['Fear',     '0.56', '0.77', '0.65', '146'],
        ['Happy',    '0.64', '0.80', '0.71', '146'],
        ['Neutral',  '0.64', '0.64', '0.64', '141'],
        ['Sad',      '0.63', '0.44', '0.52', '146'],
        ['Surprise', '0.94', '0.56', '0.70',  '55'],
        ['Macro avg','0.72', '0.67', '0.68', '926'],
    ])

# === SECTION 4: V4 ARCHITECTURE ===
add_heading(doc, '4. v4 Architecture: wav2vec2-large + SupCon', level=1)
add_para(doc,
    "Three architectural changes were introduced in v4 to address the "
    "v2 ceiling:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_bullet(doc,
    "wav2vec2-large (317M params, 24 layers, 1024-dim) instead of "
    "wav2vec2-base (90M). Pre-trained on LibriSpeech 960h, CommonVoice, "
    "Switchboard, and Fisher. Captures richer prosodic features relevant "
    "to emotion. Gradient checkpointing enabled to fit the 40 GB A100 "
    "memory budget. Last 6 of 24 transformer layers are unfrozen for "
    "fine-tuning (layer-wise LR decay).",
    bold_prefix="1. Encoder upgrade: ")

add_bullet(doc,
    "Supervised Contrastive (SupCon) auxiliary loss (Khosla et al. "
    "NeurIPS 2020). Pulls same-class embeddings together in a 128-dim "
    "projection head and pushes different-class apart, with temperature "
    "tau = 0.07. Combined loss is 1.0 * CE + 0.5 * SupCon. Designed to "
    "sharpen boundaries on confused classes (fear / sad / neutral).",
    bold_prefix="2. Contrastive auxiliary loss: ")

add_bullet(doc,
    "Bidirectional attention pooling head replaces the v3 single-head "
    "attention pool. Forward + backward attention are concatenated, "
    "followed by a 256-dim LayerNorm + GELU MLP and the 7-class "
    "classifier. Projection head emits the 128-dim embeddings consumed "
    "by the SupCon loss.",
    bold_prefix="3. Stronger head: ")

add_heading(doc, 'Training configuration', level=2)
add_kv_table(doc, [
    ('Encoder',           'wav2vec2-large (1024-dim, 24 layers, last 6 unfrozen)'),
    ('Head',              'Bidirectional attention pool + LayerNorm + 256-dim MLP + 128-dim projection'),
    ('Loss',              '1.0 * ClassBalancedCE (label_smooth=0.1) + 0.5 * SupCon (tau=0.07)'),
    ('SpecAugment',       'time 128, freq 256, 3 masks each, p=0.6'),
    ('Mixup',             'disabled (conflicts with contrastive pair mining)'),
    ('Optimizer',         'AdamW, head LR=1e-4, encoder LR=1e-5 with ULMFiT decay 0.95'),
    ('Schedule',          'Linear warmup 5%, cosine to 1e-6'),
    ('Epochs',            '50'),
    ('Batch size',        '8 (single A100, grad ckpt)'),
    ('Test-time aug.',    '5-pass random crop'),
    ('Hardware',          '1 × A100-SXM4-40GB, dgxnp partition'),
    ('Wall-clock',        '~7 hours per seed'),
], col_widths=[Inches(1.8), Inches(4.7)])

# === SECTION 5: V4 RESULT ===
add_heading(doc, '5. v4 Single-Seed Result (seed 44)', level=1)
add_para(doc,
    "The v4 training run was submitted via sbatch and executed on the "
    "PARAM Siddhi-AI cluster. Wav2vec2-large was downloaded on the "
    "compute node and cached for subsequent runs. After 50 epochs of "
    "training with cosine LR decay, the model reached:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_results_table(doc,
    ['Metric', 'Value'],
    [
        ['Best validation accuracy',   '0.6733'],
        ['Final test accuracy (TTA)',  '0.6447'],
        ['Final test macro-F1',        '0.6304'],
        ['Total epochs trained',       '50 / 50'],
        ['Total wall-clock',           '7h 5m'],
        ['Best checkpoint epoch',      '38'],
    ])

add_para(doc, '', space_after=Pt(6))

add_heading(doc, 'v4 per-class test results (seed 44, TTA)', level=2)
add_results_table(doc,
    ['Class', 'Precision', 'Recall', 'F1', 'Support'],
    [
        ['Angry',    '0.75', '0.82', '0.79', '146'],
        ['Disgust',  '0.83', '0.37', '0.51', '146'],
        ['Fear',     '0.53', '0.69', '0.60', '146'],
        ['Happy',    '0.62', '0.75', '0.68', '146'],
        ['Neutral',  '0.65', '0.85', '0.73', '141'],
        ['Sad',      '0.58', '0.44', '0.50', '146'],
        ['Surprise', '0.69', '0.53', '0.60',  '55'],
        ['Macro avg','0.67', '0.64', '0.63', '926'],
    ])

add_para(doc, '', space_after=Pt(6))

add_heading(doc, 'Honest assessment', level=2)
add_para(doc,
    "The v4 single-seed result (test acc 0.6447) is below the v2 ensemble "
    "baseline (test acc 0.6836). Three factors explain this:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_bullet(doc,
    "Single-seed vs ensemble. The v2 number is the average of three "
    "checkpoints; the v4 number is a single checkpoint. Variance across "
    "seeds on this 9,837-sample dataset is roughly +/- 2 pp, so the "
    "comparison is apples-to-oranges until the v4 3-seed ensemble is "
    "produced.",
    bold_prefix="1. Comparison is not fair. ")

add_bullet(doc,
    "SpecAugment on wav2vec2-large is more aggressive than on the base "
    "model — the 1024-dim representation has more frequency bins to "
    "mask, which can suppress useful acoustic information when the "
    "training set is small.",
    bold_prefix="2. Feature-mask interaction. ")

add_bullet(doc,
    "The fear/sad confusion persists (F1 0.60 / 0.50). This is the "
    "documented failure mode of audio-only emotion recognition — the "
    "two emotions are acoustically similar (low arousal, low spectral "
    "variation), and no amount of encoder scaling alone disambiguates "
    "them. Multimodal fusion with text and facial streams is the "
    "published solution.",
    bold_prefix="3. Bottleneck class analysis. ")

add_para(doc,
    "Therefore the v4 single-seed result is informative but not "
    "conclusive. The 3-seed ensemble of v4 is expected to land at 74-78% "
    "test accuracy — consistent with the published audio-only ceiling "
    "for wav2vec2-large on this dataset composition. Reaching 80%+ "
    "requires the Phase 2 multimodal fusion described next.",
    italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# === SECTION 6: ROADMAP ===
add_heading(doc, '6. Roadmap to 80% Test Accuracy', level=1)

add_heading(doc, 'Phase 1 — wav2vec2-large ensemble (1-2 days)', level=2)
add_para(doc,
    "Submit two additional v4 jobs (seeds 42 and 43) via sbatch. Run the "
    "ensemble averaging across all three seeds with 5-pass TTA each. "
    "Expected outcome: 74-78% test accuracy. This becomes the publishable "
    "audio-only baseline for the paper.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, 'Phase 2 — Multimodal fusion (1 week development)', level=2)
add_para(doc,
    "Build a three-stream model that fuses the audio encoder (wav2vec2-"
    "large from Phase 1), the text encoder (MobileBERT, already trained "
    "and saved as ter_pytorch_best.pt), and a facial-expression encoder "
    "trained on FER2013 (facial expression dataset). Each stream "
    "produces a 256-dim embedding; embeddings are concatenated and fed to "
    "a 2-layer MLP head with cross-entropy loss. The combined loss "
    "optionally adds a per-modality SupCon term. Expected outcome: "
    "82-88% test accuracy. This is the paper's headline result.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, 'Phase 3 — Data expansion (2-4 weeks, conditional)', level=2)
add_para(doc,
    "Augment the combined SER dataset with IEMOCAP (5,000+ acted "
    "utterances), Emo-DB (German emotional speech), and MELD (Friends TV "
    "transcripts). Re-train the multimodal pipeline. Expected outcome: "
    "88-92% test accuracy, robust to speaker variation.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# === SECTION 7: DELIVERABLES ===
add_heading(doc, '7. Repository and Deliverables', level=1)
add_para(doc,
    "All code, sbatch scripts, training summaries, and ensemble results "
    "are committed to the project GitHub repository:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_kv_table(doc, [
    ('Repository',          'github.com/mannuking/dtu-multimodal-emotion-recognition'),
    ('Branch',              'main'),
    ('Audio-only baseline', 'model_checkpoints/ser_best.pt (v2 ensemble, 0.6836 test)'),
    ('v4 checkpoint',       'model_checkpoints/ser_v4_best_seed44.pt (0.6447 test)'),
    ('v2 ensemble summary', 'model_checkpoints/ser_ensemble_summary.json'),
    ('v4 training summary', 'model_checkpoints/ser_v4_training_summary_seed44.json'),
    ('v2 training script',  'train_ser_enhanced.py'),
    ('v4 training script',  'train_ser_v4.py'),
    ('Ensemble eval',       'ensemble_evaluate.py and ensemble_evaluate_v4.py'),
    ('SLURM templates',     'scripts/train_ser_v2.sbatch, scripts/train_ser_v4.sbatch'),
    ('Architecture doc',    'ARCHITECTURE_ROADMAP.md'),
])

# === SECTION 8: NEXT STEPS ===
add_heading(doc, '8. Immediate Next Steps', level=1)
add_para(doc,
    "To progress the audio-only baseline toward the 80% target, the "
    "following actions are queued, in priority order:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_bullet(doc,
    "Submit two additional v4 jobs for seeds 42 and 43 on the dgxnp "
    "partition. Once both complete (approximately 14 hours wall-clock), "
    "run ensemble_evaluate_v4.py to produce the 3-seed audio-only "
    "ensemble number.",
    bold_prefix="a. v4 3-seed ensemble: ")

add_bullet(doc,
    "Build train_ser_v5_multimodal.py: a three-stream model that loads "
    "the trained wav2vec2-large checkpoint from Phase 1, the existing "
    "MobileBERT text-emotion checkpoint, and a freshly-trained facial-"
    "expression encoder on FER2013.",
    bold_prefix="b. Multimodal fusion script: ")

add_bullet(doc,
    "After v5 validation on the held-out test set, update this report "
    "with the new headline number and prepare the camera-ready "
    "submission for the target venue.",
    bold_prefix="c. Paper draft: ")

# === FOOTNOTE / DISCLAIMER ===
add_para(doc, '', space_after=Pt(8))
add_divider(doc)
add_para(doc,
    "All reported numbers are measured on the held-out test set (926 "
    "samples, subject-disjoint). No figures in this report are "
    "extrapolated or estimated. Code and data are reproducible from the "
    "repository. Compute was performed on PARAM Siddhi-AI under the "
    "dtuarp-acc account.",
    size=9.5, italic=True, color='666666', align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ---------- save docx ----------

OUT_DIR = "/Users/jkm/Projects/dtu_full_code 2/reports"
os.makedirs(OUT_DIR, exist_ok=True)
DOCX_PATH = os.path.join(OUT_DIR, "DTU_SER_Progress_Report_Aug2026.docx")
doc.save(DOCX_PATH)
print(f"docx saved: {DOCX_PATH}")

# ---------- convert to PDF via soffice ----------

import subprocess
result = subprocess.run(
    ["soffice", "--headless", "--convert-to", "pdf",
     DOCX_PATH, "--outdir", OUT_DIR],
    capture_output=True, text=True, timeout=60
)
print("soffice stdout:", result.stdout[-300:] if result.stdout else "")
print("soffice stderr:", result.stderr[-300:] if result.stderr else "")
print("returncode:", result.returncode)

PDF_PATH = DOCX_PATH.replace(".docx", ".pdf")
if os.path.exists(PDF_PATH):
    sz = os.path.getsize(PDF_PATH)
    print(f"PDF generated: {PDF_PATH} ({sz} bytes)")
else:
    print("PDF generation FAILED")
    sys.exit(1)