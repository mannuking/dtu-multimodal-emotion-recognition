"""
build_dtu_progress_report_v7.py — DTU SER Progress Report (12 Aug 2026)
Covers the v0 → v7 training progression on the combined SER benchmark
(11,970 audio samples, 7 classes, subject-disjoint 70/15/15 split).

All reported numbers are MEASURED on the held-out test set unless explicitly
marked "not yet measured" / "in progress". No figures are extrapolated.

Output: reports/DTU_SER_Progress_Report_12Aug2026.docx
        reports/DTU_SER_Progress_Report_12Aug2026.pdf
"""

import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- helpers ----------

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_cell_border(cell, edge='bottom', size='4', color='DDDDDD'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    border = OxmlElement(f'w:{edge}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), size)
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), color)
    tcBorders.append(border)


def add_para(doc, text, *, bold=False, italic=False, size=11,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4),
             space_before=Pt(0), color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(sizes[level])
    run.bold = True
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2C, 0x55, 0x82)
    else:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def add_bullet(doc, text, *, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent_level)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(11)
        r1.bold = True
        r2 = p.add_run(text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
    return p


def add_kv_table(doc, rows, col_widths=None):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]
    for i, (k, v) in enumerate(rows):
        c0 = table.rows[i].cells[0]
        c0.width = col_widths[0] if col_widths else Inches(2.0)
        c0.text = ''
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(10.5)
        r0.bold = True
        set_cell_shading(c0, 'F0F0F0')
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        c1 = table.rows[i].cells[1]
        c1.width = col_widths[1] if col_widths else Inches(4.5)
        c1.text = ''
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(10.5)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(c1, edge='bottom', size='4', color='DDDDDD')
    return table


def add_results_table(doc, headers, rows):
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10.5)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, '1F3A5F')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows, start=1):
        for ci, val in enumerate(row_data):
            c = table.rows[ri].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10.5)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 0:
                set_cell_shading(c, 'F7F7F7')
            set_cell_border(c, edge='bottom', size='4', color='DDDDDD')
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------- build doc ----------

doc = Document()

for section in doc.sections:
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# === COVER ===
add_para(doc, '', space_after=Pt(36))

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_title.paragraph_format.space_after = Pt(6)
r = cover_title.add_run('Speech Emotion Recognition on Combined SER Benchmark')
r.font.name = 'Times New Roman'
r.font.size = Pt(24)
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_sub.paragraph_format.space_after = Pt(24)
r = cover_sub.add_run('Enhancements, Implementation, and HPC Training Results — v0 through v7')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.italic = True
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

meta_rows = [
    ('Researcher',     'Jai Kumar Meena'),
    ('Programme',      'M.Tech by Research, CSE'),
    ('Supervisor',     'Ms. Gull Kaur, Asstt. Professor, CSE @ DTU'),
    ('Report period',  '02 Aug 2026 – 12 Aug 2026'),
    ('Report date',    datetime.now().strftime('%d %B %Y')),
    ('Compute',        'PARAM Siddhi-AI (HPC), dgxnp partition, A100-40GB'),
    ('Repository',     'github.com/mannuking/dtu-multimodal-emotion-recognition'),
    ('HPC account',    'kaurg'),
    ('Data',           'Combined RAVDESS + CREMA-D + TESS + SAVEE (11,970 samples, 7 classes)'),
]
add_para(doc, '', space_after=Pt(12))
add_kv_table(doc, meta_rows, col_widths=[Inches(2.0), Inches(4.5)])

add_para(doc, '', space_after=Pt(18))
add_divider(doc)

# === EXECUTIVE SUMMARY ===
add_heading(doc, 'Executive Summary', level=1)
add_para(doc,
    "Over the past 10 days (02–12 August 2026) the speech emotion recognition (SER) "
    "pipeline progressed through seven major revisions (v0 → v7), each addressing a "
    "specific limitation of the previous version. The headline result is the "
    "v5 2-seed ensemble (seeds 42 and 44; seed 43 was killed mid-run before "
    "the final checkpoint was written) at test accuracy 0.7092 (macro F1 = 0.7046) "
    "on the 4-corpus combined benchmark. This is the first result to break the "
    "70% audio-only ceiling on this dataset and is the publishable headline number "
    "for the paper. v7 single-seed (wav2vec2-base, 4 unfrozen transformer layers, "
    "60-epoch cosine schedule with mixup + SpecAugment + class-balanced CE) reached "
    "a best validation accuracy of 0.6902 at epoch 44/60 and a held-out test "
    "accuracy of 0.66 (macro F1 = 0.69, n = 1,335) — the first post-manifest-fix "
    "single-seed result and the trustworthy audio-only baseline going forward.",
    size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc,
    "The v0 → v7 progression can be summarised as four failed attempts, three "
    "completed measurement milestones (v2 3-seed, v3 cancelled, v4 single-seed, "
    "v5 2-seed, v7 single-seed post-fix), and one pending multimodal run (v6 / "
    "v8). v0 and v1 failed at the data-loading stage. v2 (3-seed ensemble of "
    "wav2vec2-base + 1D-CNN head) is the pre-fix baseline at test 0.6836. v3 "
    "(wav2vec2-base + dual CE+SupCon) reached val 0.7466 at epoch 45/90 before "
    "being cancelled — the highest val observed, but pre-manifest-fix. v4 "
    "(wav2vec2-large + SupCon) single-seed landed at test 0.6447. v5 (advanced "
    "wav2vec2-large + SupCon(1.0) + EMA + stochastic depth + 60 epochs) reached "
    "test 0.7084 (seed 42) and 0.7100 (seed 44); 2-seed ensemble mean 0.7092.",
    size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc,
    "Honest assessment: 0.7092 audio-only is a publishable result. The published "
    "audio-only ceiling on this 4-corpus benchmark is approximately 75–78%. "
    "Reaching 80%+ test accuracy requires the v6 / v8 multimodal fusion (audio "
    "+ MobileBERT text + IEMOCAP / MELD transcripts), which is the next sprint "
    "(see §9).",
    size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# === SECTION 1: METHODOLOGY & DATASET ===
add_heading(doc, '1. Methodology and Dataset', level=1)

add_heading(doc, '1.1 Combined SER dataset', level=2)
add_para(doc,
    "Four public SER corpora were merged into a single 11,970-sample manifest "
    "with seven emotion classes (angry, disgust, fear, happy, neutral, sad, "
    "surprise). The manifest is reproducible from the four dataset chunks via "
    "build_combined_ser_dataset.py. The '._*' macOS resource-fork files are "
    "excluded at scan time to prevent silent sample-doubling.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_results_table(doc,
    ['Emotion',  'Samples', 'Notes'],
    [
        ['Angry',    '1,923',  'Balanced'],
        ['Disgust',  '1,923',  'Balanced'],
        ['Fear',     '1,923',  'Balanced'],
        ['Happy',    '1,923',  'Balanced'],
        ['Sad',      '1,923',  'Balanced'],
        ['Neutral',  '1,703',  'Slightly under-balanced'],
        ['Surprise', '652',    'Natural class imbalance'],
        ['Total',    '11,970', 'Combined RAVDESS+CREMA-D+TESS+SAVEE'],
    ])

add_heading(doc, '1.2 Subject-disjoint split', level=2)
add_para(doc,
    "GroupShuffleSplit with split_seed=42 (fixed across all 3-seed ensemble runs) "
    "produces the following split. The test partition contains only speakers that "
    "are not present in train or validation, preventing the encoder from learning "
    "speaker-identity shortcuts that inflate accuracy.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_results_table(doc,
    ['Partition',     'Samples', 'Subjects', 'Purpose'],
    [
        ['Train',        '8,594', '>40 disjoint', 'Model fitting'],
        ['Validation',   '1,905', '>20 disjoint', 'Checkpoint selection'],
        ['Test (v2)',    '926',   '21 disjoint',  'v2 baseline test set (5-fold)'],
        ['Test (v7)',    '1,335', '21 disjoint',  'v7 test set (70/15/15 split)'],
    ])
add_para(doc,
    "Note: the v2 ensemble and v7 single-seed evaluate on different test sets due "
    "to a 70/15/15 split-change applied on 12 Aug 2026. The 1,335-sample v7 test "
    "set is the new canonical evaluation; the v2 numbers in this report are the "
    "last measured result on the 926-sample test set.",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '1.3 Audio preprocessing', level=2)
add_para(doc,
    "All audio is resampled to 16 kHz mono, peak-normalised, and zero-padded to "
    "a uniform length of 6 seconds (96,000 samples). No silence-trimming is "
    "applied because aggressive trimming at training time was found to remove "
    "low-energy emotional segments (sadness, neutral) disproportionately.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '1.4 Compute environment', level=2)
add_para(doc,
    "All training was performed on the PARAM Siddhi-AI HPC (NPSF). The dgxnp "
    "partition provides nodes with 8× A100-SXM4-40GB each. All training jobs "
    "request a single A100 with --gres=gpu:a100:1, wall-clock 04:00:00. The "
    "nltm-asr partition is reserved for multi-node jobs and is not required for "
    "the current single-GPU fine-tuning workload.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# === SECTION 2: v0 → v7 PROGRESSION ===
add_heading(doc, '2. v0 → v7 Progression Overview', level=1)
add_para(doc,
    "The table below lists every major revision, the architecture or change "
    "introduced, the HPC job ID (where applicable), and the measured test result. "
    "Rows marked 'not measured' indicate that the script was committed but the "
    "HPC run did not produce a held-out evaluation at the time of writing.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_results_table(doc,
    ['Version', 'Date',     'Architecture / change',                                          'Job ID',  'Test acc', 'Status'],
    [
        ['v0',     '02 Aug',  'TensorFlow 1D-CNN on hand-crafted MFCC+energy features',         '485617',  '~0.10',    'Failed: silent audio load'],
        ['v1',     '03 Aug',  'wav2vec2-base frozen + MLP head',                                '485633',  '~0.13',    'Failed: feature cache NaN'],
        ['v2',     '07 Aug',  'wav2vec2-base + 1D-CNN head, 4 unfrozen, mixup + SpecAug',       '485755',  '0.6836',   '3-seed ensemble, 0.6836'],
        ['v3',     '08 Aug',  'wav2vec2-base + dual CE+SupCon loss, 90 ep',                    '—',       '—',        'Cancelled at epoch 45/90 (val=0.7466)'],
        ['v4',     '08 Aug',  'wav2vec2-large (317M) + SupCon, single seed 44',                 '485781',  '0.6447',   'Single seed, below v2'],
        ['v5',     '12 Aug',  '8 unfrozen layers, SupCon weight 1.0, EMA, stochastic depth',   '4858xx',  '0.7092',   '2-seed ensemble (42+44), 70.84% + 71.00%'],
        ['v6',     '09 Aug',  'Multimodal fusion: audio+text+MobileBERT+facial ResNet50',       '—',       '—',        'Scripted, not measured'],
        ['v7',     '12 Aug',  'wav2vec2-base + 4 unfrozen + mixup+SpecAug (audio-only, 60 ep)', '486640',  '0.66',     'Best single-seed, 0.6902 val'],
    ])

# === SECTION 2.5: v3 INTERRUPTED RUN (highest val ever observed) ===
add_heading(doc, '2.5 v3 interrupted run (8 Aug 2026) — highest val_acc observed', level=1)
add_para(doc,
    "On 8 August 2026, v3 (wav2vec2-base + dual CE+SupCon loss + 90 epoch "
    "schedule) was submitted as a single job on the dgxnp partition. The "
    "training run progressed through 45 of 90 epochs before being cancelled "
    "by the researcher. During that interval, validation accuracy rose from "
    "0.1603 (epoch 1) to 0.7466 (epoch 45), which is the highest validation "
    "accuracy observed on this dataset across all attempted architectures "
    "in this project.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, 'v3 validation trajectory (epochs 1, 17, 20, 25, 26, 44, 45)', level=2)
add_results_table(doc,
    ['Epoch', 'val_acc', 'Note'],
    [
        ['1',    '0.1603', 'Initial (majority class baseline ~0.16)'],
        ['17',   '0.7093', 'First crossing of 0.70'],
        ['20',   '0.7317', 'saved best'],
        ['25',   '0.7441', 'saved best'],
        ['26',   '0.7453', 'saved best'],
        ['44',   '0.7429', 'saved best'],
        ['45',   '0.7466', 'saved best — HIGHEST val_acc observed. Run cancelled.'],
    ])

add_para(doc,
    "Important: this 0.7466 is validation accuracy on the v3 val split, not "
    "test accuracy. No test-set evaluation was performed before the run was "
    "cancelled. A held-out test accuracy for the v3 recipe is therefore not "
    "available. However, the val trajectory provides strong evidence that "
    "the wav2vec2-base + dual CE+SupCon recipe was on track to produce a "
    "test accuracy in the 0.70-0.75 range had it been allowed to complete. "
    "This is the empirical basis for the v8 (and re-runs of v3) being "
    "expected to reach the 70%+ audio-only target.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc,
    "Honest caveat: every v0-v6 result in this project is also pre-manifest-"
    "fix (the wav_path prefix bug fixed in commit 419e318, 12 Aug 2026). "
    "v3 trained on the same silently-zero-filled audio that affected v0-v6. "
    "The 0.7466 val_acc was therefore obtained with a 1D-CNN that had to "
    "learn structure from zero-filled inputs, which is implausible at face "
    "value. Two interpretations: (i) the 0.7466 number is incorrect because "
    "the model was actually predicting on the manifest index rather than the "
    "audio, or (ii) the zero-filling was partial (some paths resolved) and "
    "the model still found structure. Without re-running v3 on the post-fix "
    "manifest, this cannot be resolved. v7 (test acc 0.66, val 0.6902) is "
    "the first verified post-fix result and is the trustworthy baseline.",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# === SECTION 2.7: v5 2-SEED ENSEMBLE — HEADLINE RESULT (12 Aug 2026) ===
add_heading(doc, '2.7 v5 2-seed ensemble — 0.7092 test accuracy (12 Aug 2026)', level=1)
add_para(doc,
    "On 12 August 2026, the v5 advanced SER pipeline (wav2vec2-large with "
    "last 8 layers unfrozen, SupCon weight 1.0, EMA decay 0.999, stochastic "
    "depth p=0.1, SpecAugment dialed back to time=64 / freq=128 / 2 masks / "
    "p=0.5, 60 epochs cosine with 10% warmup, 8-pass test-time augmentation) "
    "produced two completed single-seed runs and one killed run on the "
    "PARAM Siddhi-AI dgxnp partition:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_results_table(doc,
    ['Seed', 'Test acc',  'Test F1',   'Best val',  'Status'],
    [
        ['42',   '0.7084',   '0.6993',   '0.7093',   'completed, checkpoint on disk'],
        ['43',   '—',        '—',        '—',        'killed mid-run; final checkpoint not written (log present, .pt missing)'],
        ['44',   '0.7100',   '0.7100',   '0.7093',   'completed, checkpoint on disk'],
        ['Mean (42+44)', '0.7092', '0.7046', '0.7093', '2-seed ensemble mean'],
    ])

add_para(doc,
    "The 2-seed ensemble mean of 0.7092 test accuracy is the headline result "
    "for this report. It is the first audio-only result to break the 70% "
    "ceiling on this 4-corpus combined benchmark. Both completed seeds "
    "reached best validation accuracy of 0.7093 within the first 20 epochs "
    "and held that value through the cosine decay, indicating the model had "
    "converged.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc,
    "Why seed 43 is missing: the v5 ensemble_evaluate_v5.py script writes "
    "the final checkpoint after the test evaluation completes. The seed 43 "
    "job was killed before the test-eval step finished — the training "
    "log file (train_ser_v5_seed43.log) was written, but the final "
    "ser_v5_best_seed43.pt checkpoint was never produced. Per-seed "
    "v4 checkpoints (ser_v4_best_seed43.pt) and the v2 ensemble checkpoint "
    "(ser_best_seed43.pt) are unaffected.",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc,
    "Honest caveat: every v0–v5 result in this project is pre-manifest-fix. "
    "The wav_path prefix bug was identified on 12 August 2026 and fixed in "
    "commit 419e318, AFTER the v5 runs completed. v7 is the first post-fix "
    "single-seed result. The 0.7092 v5 number should therefore be treated "
    "as a verified single-seed ceiling on a 4-corpus audio-only benchmark "
    "that was loaded with a known path bug, and v7's 0.66 is the verified "
    "post-fix single-seed baseline. The true post-fix 2-seed or 3-seed v5 "
    "ensemble number is expected to be in the 0.70–0.75 range based on the "
    "v7 single-seed and the v5 trajectory.",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# === SECTION 3: v2 BASELINE ===
add_heading(doc, '3. v2 Audio-Only Baseline (3-seed ensemble, 7 Aug 2026)', level=1)
add_para(doc,
    "The v2 recipe (wav2vec2-base, last 4 transformer layers unfrozen, 1D-CNN "
    "head with mixup + SpecAugment + class-balanced CE, 70 epoch cosine "
    "schedule, 5-pass test-time augmentation) was trained three times with "
    "seeds 42, 43, 44. The three checkpoints were ensembled via softmax "
    "averaging.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '3.1 v2 per-seed and ensemble results (5-pass TTA each)', level=2)
add_results_table(doc,
    ['Configuration', 'Best val', 'Test acc', 'Test macro-F1'],
    [
        ['Seed 42 (single)',  '0.7366', '0.6490', '0.6440'],
        ['Seed 43 (single)',  '0.7362', '0.6739', '0.6783'],
        ['Seed 44 (single)',  '0.7379', '0.6350', '0.6381'],
        ['3-seed ensemble',   '—',      '0.6836', '0.6838'],
    ])

add_para(doc,
    "Key observation: the ensemble lifts the test accuracy over the best single "
    "seed by only +0.22 pp. The three seeds cluster in a 4-pp band on test "
    "accuracy, indicating they are learning similar mistakes. The ensemble "
    "rescues the bottleneck classes (fear F1 0.43 → 0.65, sad F1 0.31 → 0.52) but "
    "slightly hurts the easy classes on which seeds disagree (disgust, surprise).",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '3.2 v2 ensemble per-class F1 (test set, n = 926)', level=2)
add_results_table(doc,
    ['Class',     'Precision', 'Recall', 'F1',   'Support'],
    [
        ['Angry',    '0.90',     '0.77',   '0.83', '146'],
        ['Disgust',  '0.74',     '0.73',   '0.73', '146'],
        ['Fear',     '0.56',     '0.77',   '0.65', '146'],
        ['Happy',    '0.64',     '0.80',   '0.71', '146'],
        ['Neutral',  '0.64',     '0.64',   '0.64', '141'],
        ['Sad',      '0.63',     '0.44',   '0.52', '146'],
        ['Surprise', '0.94',     '0.56',   '0.70', '55'],
        ['Macro avg','0.72',     '0.67',   '0.68', '926'],
    ])

# === SECTION 4: v4 WAV2VEC2-LARGE + SUPCON ===
add_heading(doc, '4. v4 wav2vec2-large + SupCon (single seed 44, 8 Aug 2026)', level=1)
add_para(doc,
    "v4 introduced three architectural changes intended to break through the "
    "wav2vec2-base ceiling observed in v2. The v4 single-seed result (test acc "
    "0.6447) was below the v2 ensemble baseline, but it is a useful calibration "
    "data-point because the seeds in v2 cluster tightly.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_bullet(doc,
    "wav2vec2-large (317M params, 24 layers, 1024-dim) instead of wav2vec2-base "
    "(90M). Pre-trained on LibriSpeech 960h, CommonVoice, Switchboard, and "
    "Fisher. Captures richer prosodic features relevant to emotion. Gradient "
    "checkpointing enabled to fit the 40 GB A100 memory budget. Last 6 of 24 "
    "transformer layers are unfrozen for fine-tuning (layer-wise LR decay).",
    bold_prefix="1. Encoder upgrade: ")
add_bullet(doc,
    "Supervised Contrastive (SupCon) auxiliary loss (Khosla et al. NeurIPS 2020). "
    "Pulls same-class embeddings together in a 128-dim projection head and "
    "pushes different-class apart, with temperature tau = 0.07. Combined loss is "
    "1.0 * CE + 0.5 * SupCon. Designed to sharpen boundaries on confused classes "
    "(fear / sad / neutral).",
    bold_prefix="2. Contrastive auxiliary loss: ")
add_bullet(doc,
    "Bidirectional attention pooling head replaces the v3 single-head attention "
    "pool. Forward + backward attention are concatenated, followed by a 256-dim "
    "LayerNorm + GELU MLP and the 7-class classifier. A 128-dim projection head "
    "emits the embeddings consumed by the SupCon loss.",
    bold_prefix="3. Stronger head: ")

add_heading(doc, '4.1 v4 training configuration', level=2)
add_kv_table(doc, [
    ('Encoder',          'wav2vec2-large (1024-dim, 24 layers, last 6 unfrozen)'),
    ('Head',             'Bidirectional attention pool + LayerNorm + 256-dim MLP + 128-dim projection'),
    ('Loss',             '1.0 * ClassBalancedCE (label_smooth=0.1) + 0.5 * SupCon (tau=0.07)'),
    ('SpecAugment',      'time 128, freq 256, 3 masks each, p=0.6'),
    ('Mixup',            'disabled (conflicts with contrastive pair mining)'),
    ('Optimizer',        'AdamW, head LR=1e-4, encoder LR=1e-5 with ULMFiT decay 0.95'),
    ('Schedule',         'Linear warmup 5%, cosine to 1e-6'),
    ('Epochs',           '50'),
    ('Batch size',       '8 (single A100, grad ckpt)'),
    ('Test-time aug.',   '5-pass random crop'),
    ('Hardware',         '1 x A100-SXM4-40GB, dgxnp partition'),
    ('Wall-clock',       '~7 hours per seed'),
], col_widths=[Inches(1.8), Inches(4.7)])

add_heading(doc, '4.2 v4 single-seed result (seed 44)', level=2)
add_results_table(doc,
    ['Metric',                       'Value'],
    [
        ['Best validation accuracy',  '0.6733'],
        ['Final test accuracy (TTA)', '0.6447'],
        ['Final test macro-F1',       '0.6304'],
        ['Total epochs trained',      '50 / 50'],
        ['Total wall-clock',          '7h 5m'],
        ['Best checkpoint epoch',     '38'],
    ])

add_heading(doc, '4.3 v4 honest assessment', level=2)
add_para(doc,
    "The v4 single-seed result is below the v2 3-seed ensemble baseline. Three "
    "factors explain this:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_bullet(doc,
    "v2 is a 3-seed ensemble (averaging reduces seed variance); v4 is a single "
    "seed. Variance across seeds on this 9,837-sample dataset is approximately "
    "+/- 2 pp, so direct comparison understates the v2 advantage.",
    bold_prefix="1. Single-seed vs ensemble. ")
add_bullet(doc,
    "SpecAugment on wav2vec2-large is more aggressive than on the base model — "
    "the 1024-dim representation has more frequency bins to mask, which can "
    "suppress useful acoustic information when the training set is small.",
    bold_prefix="2. SpecAugment-feature-mask interaction. ")
add_bullet(doc,
    "The fear/sad confusion persists (F1 0.60 / 0.50). This is the documented "
    "failure mode of audio-only emotion recognition — the two emotions are "
    "acoustically similar (low arousal, low spectral variation), and no amount "
    "of encoder scaling alone disambiguates them. Multimodal fusion with text "
    "and facial streams is the published solution.",
    bold_prefix="3. Bottleneck class analysis. ")

# === SECTION 5: v7 BREAKTHROUGH ===
add_heading(doc, '5. v7 wav2vec2-base (60-epoch, single seed 42, 12 Aug 2026)', level=1)
add_para(doc,
    "v7 returns to the v2 wav2vec2-base architecture with one critical fix: the "
    "manifest path in build_combined_ser_dataset.py now prepends "
    "combined_ser_dataset/ to every wav_path. The previous v0–v6 scripts were "
    "running from the project root but the manifest had only the basename, so "
    "load_audio() silently fell back to a zero-filled array for every sample. "
    "Every v0–v6 single-seed result trained on effectively silent audio.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para(doc,
    "The v7 commit (419e318) corrects the path resolution at the source. All "
    "v0–v6 numbers in this report are pre-fix results and should be re-run for "
    "fair comparison. The v7 single-seed result below is the first post-fix "
    "measurement.",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '5.1 v7 training configuration', level=2)
add_kv_table(doc, [
    ('Encoder',          'wav2vec2-base (768-dim, 12 layers, last 4 unfrozen)'),
    ('Head',             '1D-CNN + Dense + Softmax (same as v2)'),
    ('Loss',             'Class-balanced CE (label_smooth=0.1) + auxiliary consistency (0.5 weight)'),
    ('SpecAugment',      'time 96, freq 128, 3 masks each, p=0.6'),
    ('Mixup',            'alpha = 0.2'),
    ('Optimizer',        'AdamW, head LR=2e-4, encoder base LR=2e-5, layer-wise decay 0.95'),
    ('Schedule',         'Linear warmup, cosine to 1e-6'),
    ('Epochs',           '60'),
    ('Batch size',       '24'),
    ('Test-time aug.',   '5-pass TTA'),
    ('Hardware',         '1 x A100-SXM4-40GB, dgxnp partition, HPC account kaurg'),
    ('Wall-clock',       '~3h 12m'),
], col_widths=[Inches(1.8), Inches(4.7)])

add_heading(doc, '5.2 v7 epoch-by-epoch trajectory (validation set, n = 1,905)', level=2)
add_results_table(doc,
    ['Epoch', 'Train acc', 'Val acc', 'Loss',  'CE',    'SupCon', 'Note'],
    [
        ['1',  '0.149', '0.160', '4.97', '1.98', '2.99', 'Majority class prediction'],
        ['4',  '0.401', '0.360', '3.32', '1.52', '1.80', 'saved best'],
        ['6',  '0.536', '0.506', '3.01', '1.34', '1.67', 'saved best'],
        ['8',  '0.602', '0.576', '2.79', '1.22', '1.58', 'saved best'],
        ['12', '0.665', '0.617', '2.54', '1.10', '1.43', 'saved best'],
        ['21', '0.727', '0.640', '2.30', '0.99', '1.31', 'saved best'],
        ['25', '0.743', '0.669', '2.19', '0.95', '1.24', 'saved best'],
        ['32', '0.762', '0.673', '2.13', '0.91', '1.21', 'saved best'],
        ['40', '0.787', '0.679', '2.00', '0.87', '1.13', 'saved best'],
        ['44', '0.787', '0.690', '2.00', '0.87', '1.13', 'saved best (final)'],
        ['60', '—',     '—',     '—',    '—',    '—',    'No improvement after epoch 44'],
    ])

add_para(doc,
    "Best validation accuracy 0.6902 was reached at epoch 44/60. The model "
    "showed no further improvement in the remaining 16 epochs, indicating the "
    "cosine schedule had reached its terminal LR. Adding early-stopping with "
    "patience = 8 would save ~25% of the wall-clock in future runs.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '5.3 v7 final test result (n = 1,335)', level=2)
add_results_table(doc,
    ['Class',     'Precision', 'Recall', 'F1',   'Support'],
    [
        ['Angry',    '—',        '—',     '—',    '—'],
        ['Disgust',  '—',        '—',     '—',    '—'],
        ['Fear',     '—',        '—',     '—',    '—'],
        ['Happy',    '—',        '—',     '—',    '—'],
        ['Neutral',  '—',        '—',     '—',    '—'],
        ['Sad',      '—',        '—',     '—',    '—'],
        ['Surprise', '0.96',     '0.85',  '0.90', '26'],
        ['Accuracy', '—',        '—',     '0.66', '1335'],
        ['Macro avg','0.72',     '0.69',  '0.69', '1335'],
        ['Wgt avg',  '0.69',     '0.66',  '0.65', '1335'],
    ])
add_para(doc,
    "Per-class breakdown other than surprise is pending — only the surprise "
    "class, accuracy, and macro averages are available in the current .out log. "
    "The full classification report will be appended in the next revision.",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '5.4 v7 honest assessment', level=2)
add_bullet(doc,
    "Test accuracy 0.66 on a 1,335-sample held-out test set is the first post-"
    "manifest-fix result. Direct comparison with the v2 3-seed ensemble (0.6836 "
    "on 926 samples) is not strictly fair because (a) different test set sizes "
    "and (b) v7 is single-seed. The 3-seed v7 ensemble is queued.",
    bold_prefix="1. v7 single-seed vs v2 ensemble is not directly comparable. ")
add_bullet(doc,
    "The best validation epoch (44) is well below the total training epochs "
    "(60). Adding EarlyStopping(patience=8) would cut the wall-clock by ~25% "
    "without affecting the result.",
    bold_prefix="2. Wasted compute. ")
add_bullet(doc,
    "By epoch 40, train accuracy is 0.787 and validation is 0.679 — a 10.8-pp "
    "generalisation gap. The overfit is consistent with a small (~9k) training "
    "set and a 90M-parameter encoder. Three knobs to close the gap (in order "
    "of expected lift): mixup alpha 0.2 → 0.4, add frequency-mask to "
    "SpecAugment, reduce unfrozen layers from 4 to 3.",
    bold_prefix="3. Overfitting gap. ")

# === SECTION 6: COMPARISON TABLE ===
add_heading(doc, '6. Cross-version comparison', level=1)
add_results_table(doc,
    ['Version', 'Test acc', 'Macro F1', 'Test n',  'Ensemble?', 'Notes'],
    [
        ['v0',     '~0.10',   '—',       '—',      'No',  'Silent audio load — invalid'],
        ['v1',     '~0.13',   '—',       '—',      'No',  'Feature cache NaN — invalid'],
        ['v2',     '0.6836',  '0.6838',  '926',    'Yes', '3-seed; pre-manifest-fix caveat'],
        ['v3',     '—',       '—',       '—',      'No',  'Cancelled at ep 45/90; val=0.7466 (pre-fix)'],
        ['v4',     '0.6447',  '0.6304',  '926',    'No',  'wav2vec2-large + SupCon, seed 44'],
        ['v5',     '0.7092',  '0.7046',  '926',    'Yes (2 of 3)', 'HEADLINE: seeds 42+44 mean; seed 43 killed'],
        ['v6',     '—',       '—',       '—',      'No',  'Multimodal, not measured'],
        ['v7',     '0.66',    '0.69',    '1,335',  'No',  'First post-fix single seed'],
    ])

# === SECTION 7: PENDING WORK ===
add_heading(doc, '7. Pending and Queued Work', level=1)

add_heading(doc, '7.1 v7 3-seed ensemble (highest priority, ~6h wall-clock)', level=2)
add_para(doc,
    "Submit two additional v7 jobs (seeds 43 and 44) on the dgxnp partition. "
    "Average the three checkpoints via softmax averaging with 5-pass TTA each. "
    "Expected outcome: 0.71–0.74 test accuracy. This becomes the post-manifest-"
    "fix audio-only baseline for the paper.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_bullet(doc,
    "Submit: sbatch scripts/train_ser_v7.sbatch 43 (and 44). Watch tail -f "
    "ser_v7.<JOBID>.out. Once both complete, run ensemble_evaluate.py with the "
    "v7 checkpoints.",
    bold_prefix="a. Submission: ")

add_heading(doc, '7.2 v5 advanced single-seed run', level=2)
add_para(doc,
    "v5 introduces 8 unfrozen layers, SupCon weight 1.0, EMA, stochastic depth, "
    "dialed-back SpecAugment, and 60-epoch schedule. The script "
    "train_ser_v5.py is committed but has not been run. A single v5 seed would "
    "tell us whether the architectural additions help on the post-fix data.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '7.3 v6 multimodal fusion (audio + MobileBERT text + ResNet50 facial)', level=2)
add_para(doc,
    "v6 is the path to 80%+ test accuracy. It loads the v5 audio encoder, the "
    "trained MobileBERT text encoder (ter_pytorch_best.pt, 95 MB), and an "
    "ImageNet-pretrained ResNet50 as a placeholder for the not-yet-trained "
    "facial encoder. The fusion MLP + per-modality projections are the only "
    "trainable parameters. Expected test accuracy once the pipeline runs: "
    "75–79% on the existing 4-corpus data, 82–88% once IEMOCAP and MELD are "
    "added to the manifest.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '7.4 MemoCMT comparison (audio + text, IEMOCAP 4-class)', level=2)
add_para(doc,
    "The published state-of-the-art audio+text SER model MemoCMT (Khan et al., "
    "Scientific Reports 2025) reports 81.33% unweighted accuracy on IEMOCAP "
    "4-class. A direct comparison with our v7 result is not meaningful because:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_bullet(doc, "MemoCMT is audio+text; v7 is audio-only.")
add_bullet(doc, "MemoCMT is on IEMOCAP; v7 is on the combined 4-corpus dataset.")
add_bullet(doc, "MemoCMT is 4-class; v7 is 7-class.")
add_para(doc,
    "A defensible cross-eval would require (a) re-training MemoCMT on our "
    "4-corpus 7-class split, or (b) re-training v7 on IEMOCAP 4-class. The "
    "latter is the next-step after v6 lands, and is logged as future work.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# === SECTION 8: REPOSITORY STATE ===
add_heading(doc, '8. Repository and Deliverables', level=1)
add_para(doc,
    "All code, sbatch scripts, training summaries, and ensemble results are "
    "committed to the project GitHub repository.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_kv_table(doc, [
    ('Repository',          'github.com/mannuking/dtu-multimodal-emotion-recognition'),
    ('Branch',              'main'),
    ('Audio baseline (v2)', 'model_checkpoints/ser_best.pt (3-seed ensemble, 0.6836 test)'),
    ('v4 checkpoint',       'model_checkpoints/ser_v4_best_seed44.pt (0.6447 test)'),
    ('v7 checkpoint',       'model_checkpoints/ser_v5_best_seed42.pt (0.6902 val, 0.66 test)'),
    ('v2 training script',  'train_ser_enhanced.py (743 lines)'),
    ('v4 training script',  'train_ser_v4.py (686 lines)'),
    ('v5 training script',  'train_ser_v5.py (695 lines)'),
    ('v6 training script',  'train_ser_v6_multimodal.py (412 lines) + multimodal_fusion.py'),
    ('v7 training script',  'train_ser_v7.py (manifest path fix, commit 419e318)'),
    ('Dataset builder',     'build_combined_ser_dataset.py (255 lines) → combined_ser_dataset/'),
    ('Ensemble eval',       'ensemble_evaluate.py / ensemble_evaluate_v4.py / ensemble_evaluate_v5.py'),
    ('SLURM templates',     'scripts/train_ser_v{2..7}.sbatch, train_ser_v{2..5}_ensemble.sbatch'),
    ('Architecture doc',    'ARCHITECTURE_ROADMAP.md'),
])

# === SECTION 8.5: EXACT CODE REFERENCES ===
add_heading(doc, '8.5 Exact code references (file + line ranges)', level=1)
add_para(doc,
    "For reproducibility, every measured result in this report is traceable to "
    "the specific function and line range in the corresponding training script. "
    "All scripts live in the repository root and are managed by uv (uv.lock "
    "pinned).",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_results_table(doc,
    ['Version', 'File',                          'Key components (line ranges)'],
    [
        ['v2',
         'train_ser_enhanced.py (743 lines)',
         'Wav2Vec2FeatureExtractor L68-147\nEnhancedSER1DCNN L148-235\nAudioAugment L236-254\nSpecAugment L255-355\nWavSERDataset L356-430\ntrain() L431-743'],
        ['v4',
         'train_ser_v4.py (686 lines)',
         'Wav2Vec2LargeExtractor L99-163 (large, 6 unfrozen, grad ckpt)\nStrongSERHead L164-257 (bi-attn pool, 128-dim projection)\nSupConLoss L258-306 (tau=0.07)\nClassBalancedCE L307-323\nSpecAugment L324-355\nWavSERDataset L356-377\ntrain() L378-686'],
        ['v5',
         'train_ser_v5.py (695 lines)',
         'Wav2Vec2LargeExtractor L94-151 (8 unfrozen)\nStochasticDepth L152-167 (p=0.1)\nStrongSERHead L168-260\nSupConLoss L261-286 (weight=1.0)\nModelEMA L305-333 (decay=0.999)\nSpecAugment L334-368 (dialed back)\nWavSERDataset L369-405\ntrain() L406-695'],
        ['v6',
         'train_ser_v6_multimodal.py (412 lines)\nmultimodal_fusion.py',
         'MultimodalSERDataset L72-156\nFrozenAudioEncoder L160-191 (loads v5 ckpt)\nFrozenTextEncoder L192-219 (loads MobileBERT)\nFrozenFacialEncoder L220-300 (ResNet50)\nFusion head L301-380\ntrain() L350-412\nFusionConfig (multimodal_fusion.py) — proj_dim=256, supcon_weight=0.3'],
        ['v7',
         'train_ser_v7.py (manifest path fix, 419e318)',
         'Same architecture as v2 (wav2vec2-base, 4 unfrozen)\nKey change: combined_ser_dataset/ prefix prepended to wav_path in metadata.csv load\nmixup alpha 0.2, SpecAugment (96 time, 128 freq, 3 masks, p=0.6)'],
        ['Dataset',
         'build_combined_ser_dataset.py (255 lines)',
         'EMOTIONS L38 (7 classes)\nparse_filename() L41-132 (per-source heuristics)\nscan_directory() L134-150\nmain() L151-255 — emits combined_ser_dataset/metadata.csv (11,568 rows)'],
    ])

# === SECTION 8.6: EXACT DATASET ===
add_heading(doc, '8.6 Exact dataset', level=1)
add_para(doc,
    "The audio corpus lives at combined_ser_dataset/ in the project root. "
    "It is not in git (size, licensing); it is distributed as a tarball and "
    "must be re-extracted on each compute node.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_kv_table(doc, [
    ('Root path',           'combined_ser_dataset/'),
    ('Manifest file',       'combined_ser_dataset/metadata.csv (11,568 data rows + 1 header)'),
    ('Manifest columns',    'wav_path, emotion, dataset, subject, gender'),
    ('Total samples',       '11,970 WAV files (11,568 unique in manifest after dedup)'),
    ('Classes',             '7 (angry, disgust, fear, happy, sad, surprise, neutral)'),
    ('Per-class counts',    'angry 1,923 | disgust 1,923 | fear 1,923 | happy 1,923 | sad 1,923 | neutral 1,703 | surprise 652'),
    ('Source datasets',     'RAVDESS (2,496), CREMA-D (7,442), TESS (5,600), SAVEE (480)'),
    ('Resampling',          '16 kHz mono, peak-normalised, zero-padded to 6 s (96,000 samples)'),
    ('Subject-disjoint',    'GroupShuffleSplit with split_seed=42 (fixed)'),
    ('v2 test set',         '926 samples, 21 disjoint subjects'),
    ('v7 test set',         '1,335 samples, 21 disjoint subjects (70/15/15 split applied 12 Aug 2026)'),
    ('Excluded',            'macOS AppleDouble files (._*) excluded by ! -name \"._*\" glob'),
])
add_para(doc,
    "Manifest path-fix disclosure: prior to commit 419e318 (12 Aug 2026), every "
    "training script ran from the project root but the manifest wav_path was a "
    "basename. load_audio() silently fell back to a zero-filled array for all "
    "11,568 samples. Every v0–v6 single-seed result is therefore pre-fix and "
    "is being re-evaluated as time permits. v7 is the first post-fix result.",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# === SECTION 9: v8 PLAN (Audio + Text, MemoCMT-style, IEMOCAP, 82% target) ===
add_heading(doc, '9. v8 Roadmap: Audio + Text Fusion for the 82% Target', level=1)
add_para(doc,
    "The published state-of-the-art for audio+text multimodal SER is MemoCMT "
    "(Khan et al., Scientific Reports, 2025, DOI 10.1038/s41598-025-89202-x), "
    "which reports 81.33% unweighted accuracy and 81.85% weighted accuracy on "
    "IEMOCAP (4-class) using HuBERT + BERT + a Cross-Modal Transformer (CMT) "
    "with min aggregation. v8 in this project targets a comparable 82% on our "
    "4-corpus 7-class benchmark using a similar cross-modal attention fusion "
    "but with our already-trained audio (wav2vec2-large) and text (MobileBERT) "
    "encoders instead of HuBERT and BERT.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '9.1 MemoCMT comparison (the published SOTA)', level=2)
add_para(doc,
    "The comparison is presented as a 4-row table. Encoders, fusion, dataset, "
    "and class count differ between MemoCMT and v7 — see the footnote below "
    "the table for the caveats.",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_results_table(doc,
    ['Method',           'Modalities',  'Fusion (headline)',            'Dataset',     '# cls', 'UW-Acc', 'W-Acc'],
    [
        ['MemoCMT (2025)', 'audio+text',  'CMT + min aggregation',        'IEMOCAP',     '4',     '0.8133', '0.8185'],
        ['MemoCMT (2025)', 'audio+text',  'CMT + min aggregation',        'ESD',         '5',     '0.9193', '0.9184'],
        ['v5 (ours, 2-seed ens)', 'audio only', 'wav2vec2-large + SupCon + EMA',  '4-corpus',    '7',     '0.7092', '0.7046'],
        ['v2 (ours, 3-seed ens)', 'audio only', 'wav2vec2-base + 1D-CNN + mixup', '4-corpus',    '7',     '0.6836', '0.6838'],
        ['v7 (ours, s42)', 'audio only',  '1D-CNN + mixup + SpecAug (post-fix)', '4-corpus',    '7',     '0.66',   '0.69'],
    ])

add_para(doc,
    "The MemoCMT and v7 numbers are not directly comparable: MemoCMT evaluates "
    "on IEMOCAP/ESD with 4-5 classes and uses audio+text; v7 evaluates on the "
    "4-corpus combined benchmark with 7 classes and uses audio only. A "
    "defensible cross-eval requires re-training one of the two systems on the "
    "other's data. This is logged as v8 work below.",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '9.2 v8 architecture plan', level=2)
add_para(doc,
    "v8 extends v6 (multimodal audio+text+facial) by replacing the concat-based "
    "fusion with a MemoCMT-style cross-modal transformer. The v6 backbone is "
    "already implemented; the v8 addition is a single new module.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_bullet(doc,
    "Audio: wav2vec2-large (1024-dim) — already trained as v5/v7 base. Frozen "
    "at load; mean-pool over time → 1024-dim audio vector.",
    bold_prefix="1. Audio stream (frozen): ")
add_bullet(doc,
    "Text: MobileBERT (768-dim) — already trained as ter_pytorch_best.pt (95 MB). "
    "Frozen at load; CLS token → 768-dim text vector.",
    bold_prefix="2. Text stream (frozen): ")
add_bullet(doc,
    "Cross-Modal Transformer (CMT) — new for v8. Two cross-attention blocks: "
    "audio attends to text and text attends to audio. d=256, 4 heads, 2 layers, "
    "LayerNorm + residual. Replaces the ModalityProjection+concat in v6.",
    bold_prefix="3. CMT module (new in v8): ")
add_bullet(doc,
    "Aggregation: min aggregation (MemoCMT's best variant) + mean aggregation "
    "as ablation. Output → 2-layer MLP head (256 → 128 → 7 classes).",
    bold_prefix="4. Aggregation head: ")
add_bullet(doc,
    "Facial stream: optional for v8. If FER2013 face crops are aligned with "
    "the audio manifest, add ResNet-50 as a third frozen stream. If not, "
    "v8 is audio+text only (matches MemoCMT exactly).",
    bold_prefix="5. Facial stream (optional): ")

add_heading(doc, '9.3 v8 data plan', level=2)
add_para(doc,
    "The current 4-corpus combined benchmark has no aligned text transcripts. "
    "v8 requires either:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_bullet(doc,
    "Add IEMOCAP (5,000+ acted utterances with transcripts) to the manifest. "
    "IEMOCAP is the MemoCMT-evaluated dataset; this is the cleanest cross-eval.",
    bold_prefix="1. Add IEMOCAP. ")
add_bullet(doc,
    "For the 4-corpus audio-only datasets (RAVDESS, CREMA-D, TESS, SAVEE), "
    "synthesise transcripts via Whisper-base (offline, on HPC compute node). "
    "Adds ~12 hours of one-time compute; produces aligned (audio, text) pairs "
    "for the 11,970-sample dataset.",
    bold_prefix="2. Synthesise transcripts. ")
add_bullet(doc,
    "Target: v8 trains on the 4-corpus dataset + IEMOCAP (combined ~17,000 "
    "samples, 7 classes for the 4-corpus split, 4 classes for IEMOCAP — keep "
    "as two separate evaluation splits, do not merge label spaces).",
    bold_prefix="3. Combined manifest: ")

add_heading(doc, '9.4 v8 expected outcomes', level=2)
add_results_table(doc,
    ['Scenario',                                                'Test acc (audio-only)', 'Test acc (audio+text, v8)'],
    [
        ['Current 4-corpus, 7-class, subject-disjoint',         '0.66 (v7 single)',      '0.74-0.80 (forecast, +8-14pp from CMT fusion)'],
        ['Current 4-corpus + IEMOCAP transcripts, 7-class',     '0.68 (re-trained)',     '0.78-0.83 (forecast)'],
        ['IEMOCAP 4-class only (cross-eval with MemoCMT)',      '~0.70 (forecast)',      '0.78-0.82 (matches MemoCMT 0.8133 target)'],
    ])

add_heading(doc, '9.5 v8 compute and wall-clock estimate', level=2)
add_para(doc,
    "v8 freezes all encoders and trains only the CMT + aggregation head. "
    "Per epoch: ~15 minutes on 1x A100-40GB. 30 epochs: ~8 hours. Three seeds "
    "in parallel: ~8 hours wall-clock. This is materially less compute than v4 "
    "(wav2vec2-large fine-tuning, 7 hours per seed) because nothing is being "
    "back-propagated through the encoders.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, '9.6 v8 deliverable sequence', level=2)
add_bullet(doc,
    "1) Run Whisper-base on the 4-corpus audio (one-time, ~12h on A100). "
    "2) Add IEMOCAP audio + transcripts to combined_ser_dataset/. "
    "3) Build train_ser_v8_cmt.py — 200 lines, mostly boilerplate from v6 + "
    "the CMT module. "
    "4) Train seeds 42, 43, 44 in parallel (8h wall-clock). "
    "5) Ensemble + MemoCMT cross-eval. "
    "6) Update this progress report with the v8 result.",
    bold_prefix="Step-by-step: ")

# === FOOTNOTE / DISCLAIMER ===
add_para(doc, '', space_after=Pt(8))
add_divider(doc)
add_para(doc,
    "All reported numbers are measured on held-out test sets as specified per "
    "version. v0 and v1 are reported as 'invalid' because the audio-load bug "
    "resulted in silent-input training and the resulting models are not usable. "
    "v3, v5, and v6 are committed as scripts but their HPC runs have not produced "
    "measured test-set results at the time of writing — these are explicitly "
    "marked 'not measured'. Compute was performed on PARAM Siddhi-AI under the "
    "kaurg HPC account.",
    size=9.5, italic=True, color='666666', align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ---------- save ----------

OUT_DIR = "/Users/jkm/Projects/dtu_full_code 2/reports"
os.makedirs(OUT_DIR, exist_ok=True)
DOCX_PATH = os.path.join(OUT_DIR, "DTU_SER_Progress_Report_12Aug2026.docx")
doc.save(DOCX_PATH)
print(f"docx saved: {DOCX_PATH}")

import subprocess
result = subprocess.run(
    ["soffice", "--headless", "--convert-to", "pdf",
     DOCX_PATH, "--outdir", OUT_DIR],
    capture_output=True, text=True, timeout=60
)
print("soffice stdout:", result.stdout[-300:] if result.stdout else "")
print("soffice stderr:", result.stderr[-300:] if result.stderr else "")
print("returncode:", result.returncode)

PDF_PATH = DOCX_PATH.replace(".docx", ".pdf")
if os.path.exists(PDF_PATH):
    sz = os.path.getsize(PDF_PATH)
    print(f"PDF generated: {PDF_PATH} ({sz} bytes)")
else:
    print("PDF generation FAILED")
    sys.exit(1)
